import uuid
from pathlib import Path

import fitz
from docx import Document
from fastapi import APIRouter, Depends, File, UploadFile, status
from sqlalchemy.orm import Session

from app.db.session import get_db
from app.schemas.contract import (
    ContractCreate,
    ContractResponse,
    ContractUpdate,
)
from app.services.ai_analysis import AIAnalysisService
from app.services.contract import ContractService
from app.services.contract_analysis import ContractAnalysisService


router = APIRouter(
    prefix="/contracts",
    tags=["Contracts"],
)


@router.post(
    "",
    response_model=ContractResponse,
    status_code=status.HTTP_201_CREATED,
)
def create_contract(
    contract_data: ContractCreate,
    db: Session = Depends(get_db),
):
    service = ContractService(db)

    return service.create_contract(contract_data)


@router.get(
    "",
    response_model=list[ContractResponse],
)
def get_contracts(
    db: Session = Depends(get_db),
):
    service = ContractService(db)

    return service.get_contracts()


@router.get(
    "/{contract_id}",
    response_model=ContractResponse,
)
def get_contract(
    contract_id: uuid.UUID,
    db: Session = Depends(get_db),
):
    service = ContractService(db)

    return service.get_contract(contract_id)


@router.patch(
    "/{contract_id}",
    response_model=ContractResponse,
)
def update_contract(
    contract_id: uuid.UUID,
    contract_data: ContractUpdate,
    db: Session = Depends(get_db),
):
    service = ContractService(db)

    return service.update_contract(
        contract_id,
        contract_data,
    )


@router.delete(
    "/{contract_id}",
)
def delete_contract(
    contract_id: uuid.UUID,
    db: Session = Depends(get_db),
):
    service = ContractService(db)

    return service.delete_contract(contract_id)


@router.post(
    "/{contract_id}/upload",
    status_code=status.HTTP_200_OK,
)
async def upload_contract_file(
    contract_id: uuid.UUID,
    file: UploadFile = File(...),
):
    allowed_types = {
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }

    if file.content_type not in allowed_types:
        return {
            "error": "Only PDF and DOCX files are supported."
        }

    storage_dir = Path("storage/contracts")
    storage_dir.mkdir(parents=True, exist_ok=True)

    file_path = storage_dir / f"{contract_id}_{file.filename}"

    file_content = await file.read()
    file_path.write_bytes(file_content)

    extracted_text = ""

    if file.content_type == "application/pdf":
        pdf_document = fitz.open(file_path)

        for page in pdf_document:
            extracted_text += page.get_text()

        pdf_document.close()

    elif (
        file.content_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        document = Document(file_path)

        extracted_text = "\n".join(
            paragraph.text
            for paragraph in document.paragraphs
        )

    return {
        "contract_id": str(contract_id),
        "filename": file.filename,
        "content_type": file.content_type,
        "message": "File uploaded and text extracted successfully",
        "text": extracted_text,
    }


@router.post(
    "/{contract_id}/analyze",
    status_code=status.HTTP_200_OK,
)
def analyze_contract_document(
    contract_id: uuid.UUID,
    db: Session = Depends(get_db),
):
    storage_dir = Path("storage/contracts")

    matching_files = list(
        storage_dir.glob(f"{contract_id}_*")
    )

    if not matching_files:
        return {
            "error": "Contract file not found. Please upload the contract first."
        }

    file_path = matching_files[0]

    # Extract text from the uploaded contract
    extracted_text = ""

    if file_path.suffix.lower() == ".pdf":
        pdf_document = fitz.open(file_path)

        for page in pdf_document:
            extracted_text += page.get_text()

        pdf_document.close()

    elif file_path.suffix.lower() == ".docx":
        document = Document(file_path)

        extracted_text = "\n".join(
            paragraph.text
            for paragraph in document.paragraphs
        )

    else:
        return {
            "error": "Unsupported file format."
        }

    if not extracted_text.strip():
        return {
            "error": "No text could be extracted from the contract."
        }

    # Run AI analysis using Ollama
    ai_service = AIAnalysisService(
        model="llama3.2:latest"
    )

    analysis = ai_service.analyze_contract(
        extracted_text
    )

    # Save structured AI analysis to database
    analysis_record = ContractAnalysisService.create_analysis(
        db=db,
        contract_id=contract_id,
        summary=analysis.summary,
        key_clauses=[
            clause.model_dump()
            for clause in analysis.key_clauses
        ],
        obligations=[
            obligation.model_dump()
            for obligation in analysis.obligations
        ],
        risks=[
            risk.model_dump()
            for risk in analysis.risks
        ],
        observations=[
            observation.model_dump()
            for observation in analysis.observations
        ],
        model_name="llama3.2:latest",
    )

    return {
        "contract_id": str(contract_id),
        "filename": file_path.name,
        "message": "Contract analyzed and analysis saved successfully",
        "analysis_id": str(analysis_record.id),
        "analysis": analysis.model_dump(),
    }


@router.get(
    "/{contract_id}/analysis",
)
def get_contract_analysis(
    contract_id: uuid.UUID,
    db: Session = Depends(get_db),
):
    analysis = ContractAnalysisService.get_analysis(
        db=db,
        contract_id=contract_id,
    )

    if not analysis:
        return {
            "error": "Analysis not found for this contract."
        }

    return {
        "analysis_id": str(analysis.id),
        "contract_id": str(analysis.contract_id),
        "summary": analysis.summary,
        "key_clauses": analysis.key_clauses,
        "obligations": analysis.obligations,
        "risks": analysis.risks,
        "observations": analysis.observations,
        "model_name": analysis.model_name,
        "created_at": analysis.created_at,
    }